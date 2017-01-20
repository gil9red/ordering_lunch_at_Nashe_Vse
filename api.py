#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = 'ipetrash'


from app import session
from models import *


def has_lunch_info(lunch_id):
    """
    Функция проверяет наличие LunchInfo с таким id.

    :param lunch_id: Уникальный номер меню. Соответствует id письма, из которого было получено.
    :type lunch_id: int
    """

    has_id = session.query(LunchInfo).filter(LunchInfo.lunch_id == lunch_id).exists()
    has_id = session.query(db.literal(True)).filter(has_id).scalar()
    return True == has_id


def add_lunch(id, name, weight, price, lunch_id, category):
    """
    Функция добавляет информацию о блюде меню в базу данных.

    """

    lunch = Lunch(
        id=id,
        name=name,
        weight=weight,
        price=price,
        lunch_id=lunch_id,
        category=category
    )

    session.add(lunch)
    session.commit()


def add_lunch_info(lunch_id, subject, date):
    """
    Функция добавляет информацию о меню в базу данных и возвращает объект LunchInfo.

    """

    lunch_info = LunchInfo(lunch_id=lunch_id, subject=subject, date=date)
    session.add(lunch_info)
    session.commit()

    return lunch_info


def get_lunch_info(lunch_id):
    """
    Функция возвращает объект LunchInfo.

    """

    return session.query(LunchInfo).filter(LunchInfo.lunch_id == lunch_id).one()


def get_lunchs_list(lunch_id):
    """
    Функция возвращает список объектов Lunch, относящихся к определенному меню (lunch_id).

    """

    return session.query(Lunch).filter(Lunch.lunch_id == lunch_id).order_by(Lunch.id)


def get_user_order_lunch(user_ip, lunch_id):
    """
    Функция возвращает объект UserOrderLunch.

    """

    return session.query(UserOrderLunch)\
        .filter(UserOrderLunch.user_ip == user_ip)\
        .filter(UserOrderLunch.lunch_id == lunch_id).one_or_none()


def get_or_create_new_user_order_lunch(user_ip, lunch_id, selected_id_menu=None, additionally=None):
    user_lunch = get_user_order_lunch(user_ip, lunch_id)

    if user_lunch is None:
        user_lunch = UserOrderLunch(user_ip=user_ip, name=get_host_name(user_ip),
                                    lunch_id=lunch_id, additionally=additionally)
        session.add(user_lunch)

    import datetime
    user_lunch.last_online_date = datetime.datetime.today()

    if selected_id_menu is not None:
        # Конвертируем элементы списка в строку, т.к. join работает
        # только с строковыми значениями
        user_lunch.order = ','.join([str(_) for _ in selected_id_menu])
        user_lunch.last_select_lunch_date = user_lunch.last_online_date

    if additionally is not None:
        user_lunch.additionally = additionally

    session.commit()
    return user_lunch


def get_row_lunchs_menu_for_html_render(lunch_id, user):
    """
    Функция возвращает оформленный для рендеринга список меню.

    """

    lunch_id_list = user.get_lunch_id_list()

    # Теперь нужно заполнить список для генератора шаблона
    rows = list()
    category = None

    # Запрашиваем блюда за lunch_id (последние блюда)
    for lunch in get_lunchs_list(lunch_id):
        if category != lunch.category:
            category = lunch.category

            rows.append(
                {
                    "category": category,
                    'is_item': False,  # Флаг для генератора шаблонов
                }
            )

        rows.append(
            {
                'id': lunch.id,
                'name': lunch.name,
                'weight': lunch.weight,
                'price': lunch.price,

                # Флаг для генератора шаблонов
                'is_item': True,

                # Флаг выбранного блюда. Проверка что текущие id блюда есть
                # среди выбранных пользователем
                'is_selected': False if user is None else lunch.id in lunch_id_list,
            }
        )

    return rows


def get_total_price_and_weight(rows):
    """
    Функция для подсчета и возврата общего веса и цены списка блюд.

    """

    from collections import defaultdict
    weight_number = defaultdict(int)
    price_number = defaultdict(int)

    # Список для нестандартных блюд вида "Блинчики со сгущенкой/джемом/ сметаной  100/50 гр 45/55/55 руб. "
    other_weights = list()
    other_prices = list()

    for row in rows:
        weight, weight_value = row['weight'].strip().split()
        price, price_value = row['price'].strip().split()

        try:
            weight_number[weight_value] += int(weight)
        except:
            other_weights.append(weight + ' ' + weight_value)

        try:
            price_number[price_value] += int(price)
        except:
            other_prices.append(price + ' ' + price_value)

    # Подсчет общего веса и цены
    total_weights = sorted(weight_number.items(), key=lambda x: x[1], reverse=True)
    total_weights = ['{} {}'.format(num, value) for value, num in total_weights]
    total_weights += other_weights

    total_prices = sorted(price_number.items(), key=lambda x: x[1], reverse=True)
    total_prices = ['{} {}'.format(num, value) for value, num in total_prices]
    total_prices += other_prices

    return total_weights, total_prices


def get_host_name(ip):
    """
    Функция по указанному ip возвращает имя хоста.

    """

    if ip == "127.0.0.1":
        return "localhost"

    import socket
    return socket.gethostbyaddr(ip)[0]


import logging
import os
import config


def save_attachment(msg, file_name):
    """
    Given a message, save its attachments to the specified
    download folder (default is /tmp)

    """

    for part in msg.walk():
        if part.get_content_maintype() == 'multipart':
            continue

        if part.get('Content-Disposition') is None:
            continue

        if not os.path.exists(file_name):
            with open(file_name, 'wb') as fp:
                fp.write(part.get_payload(decode=True))
        else:
            logging.debug('Lunch menu file exist: %s', file_name)

        break


def last_email_id():
    """
    Функция для получения id последнего письма.

    """

    # Для тестирования сервера без доступа к почте, имея только базу в памяти.
    # last_lunch_id заранее заполняется
    if config.debug_without_email:
        return None, None, 9999

    logging.debug('Check last email.')

    import imaplib
    connect = imaplib.IMAP4(config.smtp_server)
    connect.login(config.username, config.password)
    connect.select()

    # Если не ограничивать датой, соберет все письма и запрос будет дольше выполняться
    from datetime import date, timedelta
    today = date.today()
    week_ago = today - timedelta(weeks=1)
    since = week_ago.strftime('%d-%b-%Y')

    logging.debug('Search emails from %s, since: %s.', config.lunch_email, since)
    typ, msgnums = connect.search(None, 'HEADER From', config.lunch_email, 'SINCE', since)
    # NOTE: поиск всех писем
    # typ, msgnums = connect.search(None, 'HEADER From', config.lunch_email)
    logging.debug('Search result: %s.', msgnums[0].split())
    # TODO: письма можно получить все, а после сортировать по дате получения и уже брать последнее
    # т.к. id писем не гарантируют порядок получения -- если из inbox переместить письма в другую папку
    # и часть вернуть обратно, у них, похоже будет новый id. И так может получиться, что более старые
    # письма окажутся последними полученными от SEARCH

    id_list = msgnums[0].split()
    if not id_list:
        raise Exception('Письма не найдены.')

    last_id = id_list[-1]
    return connect, last_id, int(last_id.decode())


def save_last_lunch_menu():
    """
    Функция получает последнее письмо от указанного емейла
    и сохраняет из него меню в базу данных.

    """

    # Регулярка для поиска последовательностей пробелов: от двух подряд и более
    import re
    multi_space_pattern = re.compile(r'[ ]{2,}')

    try:
        connect = None
        connect, last_id, last_id_int = last_email_id()

        if has_lunch_info(last_id_int):
            logging.debug('Lunch menu #%s already exist!', last_id_int)
            return last_id_int

        logging.debug('Получение письма с uid: %s.', last_id_int)
        typ, data = connect.fetch(last_id, '(RFC822)')

        import email
        msg = email.message_from_bytes(data[0][1])

        # Получаем заголовок письма
        data, charset = email.header.decode_header(msg['Subject'])[0]
        subject = data.decode(charset).strip().replace('\t', ' ')
        logging.debug('Subject: "%s".', subject)

        # Получаем дату получения (отправления??) письма
        from datetime import datetime
        date_tuple = email.utils.parsedate_tz(msg['Date'])
        email_date = datetime.fromtimestamp(email.utils.mktime_tz(date_tuple))
        logging.debug("Date: %s.", email_date)

        # Добавляем информацию о последнем меню
        lunch_info = add_lunch_info(last_id_int, subject, email_date)

        try:
            file_name = 'last_lunch_menu.docx'
            save_attachment(msg, file_name)

            logging.debug('Read lunch file name: %s.', file_name)
            from docx import Document
            document = Document(file_name)

            category = None

            if document.tables:
                # Таблицы в меню дублируются
                table = document.tables[0]

                # Перебор начинается со второй строки, потому что, первая строка таблицы -- это строка "Обеденное меню"
                i = 1
                for row in table.rows[1:]:
                    name, weight, price = [multi_space_pattern.sub(' ', i.text.strip()) for i in row.cells]

                    # Заметил, что у категорий или совпадают 3 столбца, или отсутвует второй и третий столбцы
                    if name == weight == price or (not weight or not price):
                        category = name.title()
                        logging.debug(category)
                        continue

                    add_lunch(i, name, weight, price, lunch_info.lunch_id, category)
                    i += 1

                    logging.debug('{} {} {}'.format(name, weight, price))

        finally:
            if os.path.exists(file_name):
                os.remove(file_name)

    finally:
        if connect:
            connect.close()
            connect.logout()

    return last_id_int


# TODO: decline
def get_today_users():
    """
    Функция возвращает список пользователей заходивших на вебстраницу
    сегодня и смотревших последнее доступное меню

    """

    _, _, last_id_int = last_email_id()

    from datetime import datetime

    return session.query(UserOrderLunch)\
        .filter(db.func.date(UserOrderLunch.last_online_date) == datetime.today().date())\
        .filter(UserOrderLunch.lunch_id == last_id_int)\
        .all()


def get_last_menu_users():
    """
    Функция возвращает список пользователей, смотревших последнее доступное меню

    """

    _, _, last_id_int = last_email_id()

    # from datetime import datetime

    return session.query(UserOrderLunch)\
        .filter(UserOrderLunch.lunch_id == last_id_int)\
        .all()


# TODO: decline
# def get_users_by_online_date(online_date=None):
#     """
#     Функция возвращает список пользователей заходивших на вебстраницу
#     за указанную дату.
#
#     """
#
#     from datetime import datetime
#     if online_date is None:
#         online_date = datetime.today().date()
#     else:
#         online_date = datetime.strptime(online_date, '%d%m%Y').date()
#
#     return session.query(UserOrderLunch)\
#         .filter(db.func.date(UserOrderLunch.last_online_date) == online_date)\
#         .all()


def get_users_by_online_date(online_date=None):
    """
    Функция возвращает список пользователей заходивших на вебстраницу
    за указанную дату.

    """

    if online_date is None:
        from datetime import datetime
        online_date = datetime.today().date().strftime('%d%m%Y')

    return get_users_by_range_online_dates(online_date, online_date)


def get_users_by_range_online_dates(start_user_online_date, last_user_online_date, reverse=False):
    """
    Функция возвращает список пользователей за указанный диапазон дат.

    """

    from datetime import datetime
    start_user_online_date = datetime.strptime(start_user_online_date, '%d%m%Y').date()
    last_user_online_date = datetime.strptime(last_user_online_date, '%d%m%Y').date()

    column_date = db.func.date(UserOrderLunch.last_online_date)

    query = session.query(UserOrderLunch)\
        .filter(column_date >= start_user_online_date)\
        .filter(column_date <= last_user_online_date)

    if reverse:
        column_date = query.order_by(column_date.desc())
    else:
        column_date = query.order_by(column_date.asc())

    return column_date.all()


def create_url_file(file_name):
    """
    Функция по текущему адресу сервера создает .url файл.

    """

    from flask import request

    text = """\
[InternetShortcut]
URL=http://{}/
    """.format(request.host)

    with open(file_name, mode='w', encoding='ascii') as f:
        f.write(text)


def send_email(subject, text):
    """Функция отправляет письмо на указанные в конфиге почты."""

    try:
        import config

        # typical values for text_subtype are plain, html, xml
        text_subtype = 'plain'

        from email.mime.text import MIMEText
        msg = MIMEText(text, text_subtype)
        msg['Subject'] = subject
        msg['From'] = config.sender
        msg['To'] = config.to_email
        msg['Cc'] = ', '.join(config.to_cc_emails)

        # # this invokes the secure SMTP protocol (port 465, uses SSL)
        # from smtplib import SMTP_SSL as SMTP

        # use this for standard SMTP protocol   (port 25, no encryption)
        from smtplib import SMTP

        with SMTP(config.smtp_server) as smtp:
            smtp.set_debuglevel(config.debug_smtp)
            smtp.starttls()
            smtp.login(config.username, config.password)
            smtp.send_message(msg)

    except Exception as e:
        print("mail failed; {}".format(e))


if __name__ == '__main__':
    from datetime import date
    subject = 'Заказ меню за {}'.format(date.today().strftime("%d/%m/%Y"))
    text = 'Большой бургер!'

    send_email(subject, text)

    # users = get_users_by_range_online_dates('18052016', '21052016')
    # for user in users:
    #     print(user.last_online_date, user)
    #
    # print()
    # users = get_users_by_range_online_dates('18052016', '21052016', True)
    # for user in users:
    #     print(user.last_online_date, user)
    #
    # print()
    # users = get_users_by_online_date('21052016')
    # for user in users:
    #     print(user.last_online_date, user)
    #
    # print()
    # users = get_users_by_online_date()
    # for user in users:
    #     print(user.last_online_date, user)
